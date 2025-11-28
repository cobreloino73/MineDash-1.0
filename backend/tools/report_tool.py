"""
MineDash AI - Report Generator Tool
Herramienta para generar reportes profesionales en Word
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime


class ReportGenerator:
    """
    Generador de reportes profesionales en Word
    """
    
    def __init__(self, reports_dir: Path):
        """
        Inicializar Report Generator
        
        Args:
            reports_dir: Directorio donde guardar reportes
        """
        self.reports_dir = Path(reports_dir)
        self.reports_dir.mkdir(parents=True, exist_ok=True)
        
        # Colores corporativos Codelco
        self.color_primary = RGBColor(230, 57, 70)
        self.color_secondary = RGBColor(17, 138, 178)
        self.color_text = RGBColor(50, 50, 50)
    
    def generate(
        self,
        title: str,
        sections: List[Dict[str, Any]],
        format_type: str = 'docx'
    ) -> Path:
        """
        Generar reporte
        
        Args:
            title: Título del reporte
            sections: Lista de secciones
            format_type: Formato ('docx')
            
        Returns:
            Path al archivo generado
        """
        if format_type == 'docx':
            return self._generate_docx(title, sections)
        else:
            raise ValueError(f"Formato no soportado: {format_type}")
    
    def _generate_docx(self, title: str, sections: List[Dict[str, Any]]) -> Path:
        """Generar reporte en Word"""
        doc = Document()
        
        # Portada
        self._add_cover(doc, title)
        
        # Contenido
        for section in sections:
            heading = section.get('heading', 'Sin título')
            content = section.get('content', '')
            chart_path = section.get('chart_path')
            table_data = section.get('table')
            
            # Encabezado
            self._add_heading(doc, heading, level=1)
            
            # Contenido
            if content:
                self._add_paragraph(doc, content)
            
            # Gráfico
            if chart_path and Path(chart_path).exists():
                try:
                    doc.add_picture(chart_path, width=Inches(6))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except:
                    pass
            
            # Tabla
            if table_data:
                self._add_table(doc, table_data)
            
            doc.add_paragraph()
        
        # Guardar
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"reporte_{timestamp}.docx"
        filepath = self.reports_dir / filename
        
        doc.save(filepath)
        
        return filepath
    
    def _add_cover(self, doc: Document, title: str):
        """Agregar portada"""
        # Título
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(f"\n\n\n\n\n{title}")
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = self.color_primary
        
        # Subtítulo
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run('División Salvador - Codelco Chile')
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.color.rgb = self.color_secondary
        
        # Fecha
        date_para = doc.add_paragraph('\n\n\n\n\n\n\n\n')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(
            f'Generado: {datetime.now().strftime("%d-%m-%Y")}'
        )
        date_run.font.size = Pt(12)
        
        doc.add_page_break()
    
    def _add_heading(self, doc: Document, text: str, level: int = 1):
        """Agregar encabezado"""
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            if level == 1:
                run.font.size = Pt(18)
                run.font.color.rgb = self.color_primary
    
    def _add_paragraph(self, doc: Document, text: str):
        """Agregar párrafo"""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = self.color_text
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _add_table(self, doc: Document, table_data: Dict[str, Any]):
        """Agregar tabla"""
        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])
        
        if not headers or not rows:
            return
        
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Encabezados
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = str(header)
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
        
        # Filas
        for i, row_data in enumerate(rows, start=1):
            row_cells = table.rows[i].cells
            for j, cell_data in enumerate(row_data):
                row_cells[j].text = str(cell_data)
        
        doc.add_paragraph()